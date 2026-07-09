from io import BytesIO
from pathlib import Path
import uuid

import pytest
from docx import Document

from app.services.fc_doc_parser import (
    FcDocParseError,
    FcDocSizeLimitError,
    UnsupportedFcDocFormatError,
    get_fc_upload_directory,
    parse_requirement_content,
    parse_txt_or_md,
    save_fc_requirement_upload,
)


def test_parse_txt_content() -> None:
    content = b"# Login PRD\n\nUser can login with password."
    result = parse_txt_or_md(content)
    assert "Login PRD" in result
    assert "User can login" in result


def test_parse_md_content() -> None:
    content = b"## Requirement\n- Case A"
    result = parse_requirement_content(content, "requirements.md")
    assert "Requirement" in result
    assert "Case A" in result


def test_parse_docx_content() -> None:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Payment module requirement")
    document.add_paragraph("Support refund flow")
    document.save(buffer)
    content = buffer.getvalue()

    result = parse_requirement_content(content, "payment.docx")
    assert "Payment module requirement" in result
    assert "Support refund flow" in result


def test_parse_rejects_unsupported_extension() -> None:
    with pytest.raises(UnsupportedFcDocFormatError):
        parse_requirement_content(b"plain", "notes.pdf")


def test_parse_rejects_empty_txt() -> None:
    with pytest.raises(FcDocParseError):
        parse_txt_or_md(b"   \n  ")


def test_save_fc_requirement_upload_writes_file(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr("app.services.fc_doc_parser.STORAGE_DIR", tmp_path)
    project_id = uuid.uuid4()
    content = b"Requirement text"

    saved_path = save_fc_requirement_upload(project_id, "demo.txt", content)

    assert saved_path.parent == get_fc_upload_directory(project_id)
    assert saved_path.suffix == ".txt"
    assert saved_path.read_bytes() == content


def test_save_fc_requirement_upload_accepts_chinese_filename(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr("app.services.fc_doc_parser.STORAGE_DIR", tmp_path)
    project_id = uuid.uuid4()
    content = b"Requirement text"

    saved_path = save_fc_requirement_upload(project_id, "需求文档.docx", content)

    assert saved_path.suffix == ".docx"
    assert saved_path.read_bytes() == content


def test_save_fc_requirement_upload_rejects_oversized_file(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr("app.services.fc_doc_parser.STORAGE_DIR", tmp_path)

    class FakeSettings:
        fc_max_doc_size_mb = 1

    monkeypatch.setattr("app.services.fc_doc_parser.get_settings", lambda: FakeSettings())

    oversized = b"x" * (1024 * 1024 + 1)
    with pytest.raises(FcDocSizeLimitError):
        save_fc_requirement_upload(uuid.uuid4(), "large.txt", oversized)


def test_save_fc_requirement_upload_rejects_bad_extension(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr("app.services.fc_doc_parser.STORAGE_DIR", tmp_path)
    with pytest.raises(UnsupportedFcDocFormatError):
        save_fc_requirement_upload(uuid.uuid4(), "demo.pdf", b"pdf")
