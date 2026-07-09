from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient


def _create_fc_project(client: TestClient, auth_headers: dict[str, str]) -> str:
    response = client.post(
        "/api/v1/fc-projects",
        headers=auth_headers,
        json={"name": "Doc Test Project", "description": "For requirement docs"},
    )
    assert response.status_code == 201
    return response.json()["id"]


def _build_docx_bytes() -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("User login requirement")
    document.add_paragraph("Must validate password length")
    document.save(buffer)
    return buffer.getvalue()


def test_upload_requirement_doc_docx(
    client: TestClient,
    auth_headers: dict[str, str],
) -> None:
    project_id = _create_fc_project(client, auth_headers)
    upload_url = f"/api/v1/fc-projects/{project_id}/docs/upload"

    response = client.post(
        upload_url,
        headers=auth_headers,
        files={"file": ("demo-prd.docx", _build_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 201
    payload = response.json()["doc"]
    assert payload["filename"] == "demo-prd.docx"
    assert payload["file_type"] == "docx"
    assert payload["parse_status"] == "success"
    assert payload["parse_error"] is None
    assert "User login requirement" in (payload["parsed_text_preview"] or "")

    doc_id = payload["id"]

    list_response = client.get(
        f"/api/v1/fc-projects/{project_id}/docs",
        headers=auth_headers,
    )
    assert list_response.status_code == 200
    assert len(list_response.json()) == 1

    detail_response = client.get(
        f"/api/v1/fc-projects/{project_id}/docs/{doc_id}",
        headers=auth_headers,
    )
    assert detail_response.status_code == 200
    detail = detail_response.json()
    assert detail["parse_status"] == "success"
    assert "Must validate password length" in (detail["parsed_text"] or "")

    delete_response = client.delete(
        f"/api/v1/fc-projects/{project_id}/docs/{doc_id}",
        headers=auth_headers,
    )
    assert delete_response.status_code == 204


def test_upload_requirement_doc_txt(
    client: TestClient,
    auth_headers: dict[str, str],
) -> None:
    project_id = _create_fc_project(client, auth_headers)
    response = client.post(
        f"/api/v1/fc-projects/{project_id}/docs/upload",
        headers=auth_headers,
        files={"file": ("requirements.txt", b"Login requirement details", "text/plain")},
    )
    assert response.status_code == 201
    assert response.json()["doc"]["parse_status"] == "success"


def test_upload_requirement_doc_rejects_unsupported_format(
    client: TestClient,
    auth_headers: dict[str, str],
) -> None:
    project_id = _create_fc_project(client, auth_headers)
    response = client.post(
        f"/api/v1/fc-projects/{project_id}/docs/upload",
        headers=auth_headers,
        files={"file": ("notes.pdf", b"%PDF", "application/pdf")},
    )
    assert response.status_code == 400


def test_upload_requirement_doc_chinese_filename(
    client: TestClient,
    auth_headers: dict[str, str],
) -> None:
    project_id = _create_fc_project(client, auth_headers)
    response = client.post(
        f"/api/v1/fc-projects/{project_id}/docs/upload",
        headers=auth_headers,
        files={
            "file": (
                "需求文档.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201
    payload = response.json()["doc"]
    assert payload["filename"] == "需求文档.docx"
    assert payload["parse_status"] == "success"


def test_upload_requirement_doc_requires_auth(client: TestClient) -> None:
    import uuid

    project_id = uuid.uuid4()
    response = client.post(
        f"/api/v1/fc-projects/{project_id}/docs/upload",
        files={"file": ("demo.txt", b"text", "text/plain")},
    )
    assert response.status_code == 403
